from flask import Flask, render_template, request, send_file
from deep_translator import GoogleTranslator
from docx import Document
from fpdf import FPDF
import time
import io

app = Flask(__name__)

app.config['SECRET_KEY'] = 'systemctl'
app.config['DEBUG'] = False 

def enhance_horror_tone(text):
    horror_words = {
        "strange": "perturbador",
        "dark": "sombrio",
        "cold": "gélido",
        "silent": "silencioso demais",
        "alone": "totalmente só",
        "shadow": "sombra ameaçadora",
        "fear": "terror profundo",
        "death": "morte iminente",
        "creep": "arrepio na espinha"
    }
    for word, horror_word in horror_words.items():
        text = text.replace(word, horror_word)
    return text

def translate_large_text(text, source_lang='en', target_lang='pt', chunk_size=3000):
    segments = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    translated_text = ""
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    for segment in segments:
        try:
            translated_segment = translator.translate(segment)
            translated_segment = enhance_horror_tone(translated_segment)
            translated_text += translated_segment + " "
            time.sleep(1)
        except Exception as e:
            translated_text += " [Erro na tradução de um segmento] "
    return translated_text.strip()

@app.route('/', methods=['GET', 'POST'])
def translate():
    translation = ""
    if request.method == 'POST':
        original_text = request.form['text']
        source_language = request.form.get('source_language', 'en')  # Padrão: inglês
        try:
            # Traduzir do idioma selecionado para português
            translation = translate_large_text(original_text, source_lang=source_language, target_lang='pt')
        except Exception as e:
            translation = "Erro na tradução. Tente novamente mais tarde."
    return render_template('index.html', translation=translation)



@app.route('/download', methods=['POST'])
def download():
    translated_text = request.form['translated_text']
    filename = "CreepyTradutor.txt"
    file_content = f"--- CreepyTradutor ---\n\n{translated_text}"

    return send_file(
        io.BytesIO(file_content.encode('utf-8')),
        as_attachment=True,
        download_name=filename,
        mimetype='text/plain'
    )
    
@app.route('/downloadDocx', methods=['POST'])
def downloadDocx():
    translated_text = request.form['translated_text']
    filename = "CreepyTradutor.docx"

    # Criar um documento .docx usando python-docx
    doc = Document()
    doc.add_heading('--- CreepyTradutor ---', level=1)
    doc.add_paragraph(translated_text)

    # Salvar o documento em um buffer de memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Retornar o arquivo como resposta
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
@app.route('/downloadPDF', methods=['POST'])
def downloadPDF():
    translated_text = request.form['translated_text']
    filename = "CreepyTradutor.pdf"

    # Criar o PDF
    pdf = FPDF()
    pdf.add_page()

    # Adicionar uma fonte Unicode (DejaVu Sans)
    pdf.add_font('DejaVu', '', './fonts/DejaVuSans.ttf', uni=True)
    pdf.set_font('DejaVu', size=12)

    # Adicionar título
    pdf.set_font('DejaVu', size=16)
    pdf.cell(0, 10, "--- CreepyTradutor ---", ln=True, align='C')

    # Adicionar texto traduzido
    pdf.set_font('DejaVu', size=12)
    pdf.multi_cell(0, 10, translated_text)

    # Salvar o PDF em um buffer de memória
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)

    # Retornar o arquivo PDF para download
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )

if __name__ == '__main__':
   # app.run(port=5000)
    app.run(debug=True)